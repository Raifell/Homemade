from bs4 import BeautifulSoup
import requests
import random
import time
import json
import glob
import docx

url = 'https://www.olx.uz/rabota/q-python/'
heders = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/114.0.0.0 Safari/537.36 OPR/100.0.0.0',
          'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.7'
          }

def harvest(url):
    page_main = requests.get(url=url, headers=heders)
    soup = BeautifulSoup(page_main.text, features='html.parser')

    catalog = soup.find_all('div', class_='css-1sw7q4x')[:40]
    link = ['https://www.olx.uz' + x.find('a')['href'] for x in catalog]

    for x in range(len(link)):
        page_work = requests.get(url=link[x], headers=heders)
        soup = BeautifulSoup(page_work.text, features='html.parser').find('div', class_='css-1m8mzwg')
        # p = [x.text for x in soup.find_all('p')]
        # li = [x.text for x in soup.find_all('li')]
        div = [x.text for x in soup.find_all('div')]
        with open(f'text/site-{x}.txt', 'w', encoding='utf-8') as f:
            f.write('\n'.join(div))

        print('Complete {}'.format(x))
        time.sleep(random.randint(5, 15))

    count = {'sql': 0, 'docker': 0, 'django': 0, 'flask': 0, 'fastapi': 0, 'nginx': 0, 'apache': 0, 'git': 0, 'unix': 0,
             'rest': 0, 'python': 0, 'linux': 0, 'figma': 0, 'php': 0, 'soap': 0, 'debian': 0, 'java': 0, 'kafka': 0,
             'redis': 0, 'celery': 0, 'react': 0, 'html': 0, 'css': 0}

    for x in range(40):
        with open(f'text/site-{x}.txt', 'r', encoding='utf-8') as f:
            text = f.readlines()
            for i in text:
                for z in count:
                    if z in i.lower():
                        count[z] += 1

    res = {k: v for k, v in sorted(count.items(), key=lambda item: item[1], reverse=True)}
    result = json.dumps({'olx.uz': res}, indent=3)
    with open(f'result/olx.json', 'w') as f:
        f.write(result)


files = glob.glob("result/*")

doc = docx.Document()

for i in files:
    with open(i, 'r') as f:
        data_load = json.loads(f.read())
        for x in data_load:
            num = 20 if i != 'result\\olx.json' else 40
            doc.add_heading('Сайт\\Запрос - {}'.format(x), 1)
            doc.add_paragraph(f'Спрос на 10 ключевых навыков, относительно {num} запросов:')
            for y, z in enumerate(data_load[x].items()):
                doc.add_paragraph(' '.join([str(z[0]).capitalize(), ':', str(z[1])]))
                if y == 9:
                    break

doc.add_heading('Сайты вошедшие в статистику', 1)
re_files = [x for x in set(i.replace('result\\', '').replace('.json', '')[:-1] for i in files)]
for x in range(len(re_files)):
    if re_files[x] == 'ol':
        re_files[x] = 'olx'
    elif re_files[x] == 'ishko':
        re_files[x] = 'ishkop'

for i in re_files:
    doc.add_paragraph(i + ' (ru/uz)')
doc.add_paragraph('Всего страниц проинспектировано: 240')
doc.save('over/HW_27_09_2023.docx')

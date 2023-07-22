# Задание №1
# а) Создайте word файл в операционной системе, заполните его текстом «Hello Python».
# б) Прочитайте этот файл, только жирный текст в python-строку и выведите на экран.
# в) Создайте абзац с текстом и запишите это в новый word-файл, изменит шрифт и размер шрифта.

# Для выполнения задания взял набор текста из HTML модуля lorem, генератор текста на его основе https://loremipsum.io

import docx
from docx.shared import Pt


def create_start(path):
    doc = docx.Document()

    doc.add_paragraph('Homework', 'Title').alignment = 1
    p1 = doc.add_paragraph('    Lorem ipsum dolor sit amet, consectetur adipiscing elit, ')
    p1.add_run('sed do eiusmod tempor ')
    p1.add_run('incididunt').bold = True
    p1.add_run(' ut labore et dolore magna aliqua.')

    p2 = doc.add_paragraph('    Ut enim ad minim veniam, ')
    p2.add_run('quis nostrud exercitation ullamco ').bold = True
    p2.add_run('laboris nisi ut aliquip ex ea commodo consequat.')

    p3 = doc.add_paragraph('    Duis aute irure dolor in reprehenderit in voluptate velit ')
    p3.add_run('esse cillum dolore eu fugiat nulla pariatur. ').bold = True
    p3.add_run('Excepteur sint occaecat cupidatat non proident, ')
    p3.add_run('sunt in culpa qui officia').bold = True
    p3.add_run('deserunt mollit anim id est laborum.')

    doc.save(path)


def count_bold(path, new_path):
    doc = docx.Document(path)

    bolds = []
    for para in doc.paragraphs:
        for run in para.runs:
            if run.bold == True:
                bolds.append(run.text)
    create_new(bolds, new_path)


def create_new(bolds, new_path):
    doc = docx.Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(14)

    ps = doc.add_paragraph('')

    xs = ps.add_run(' '.join(bolds))
    xs.font.size = Pt(26)
    xs.font.italic = True

    doc.save(new_path)


create_start('example.docx')
count_bold('example.docx', 'example1.docx')
